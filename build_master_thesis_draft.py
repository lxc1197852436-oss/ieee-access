from pathlib import Path
import shutil

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = Path("/Users/xingchengli/Documents/New project/硕士学位论文初稿_中国数据扩展版.docx")
FIG_DIR = Path("/Users/xingchengli/Documents/New project/ieee_pkg/overleaf_ieee_access_package_20260612/figures")


def set_font(run, name="宋体", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para(paragraph, before=0, after=6, line=1.35, first_line=True, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    if align is not None:
        paragraph.alignment = align


def add_p(doc, text="", style=None, bold_prefix=None, first_line=True):
    p = doc.add_paragraph(style=style)
    set_para(p, first_line=first_line)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_h(doc, text, level=1):
    p = doc.add_heading("", level=level)
    p.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_font(run, "黑体", 16 if level == 1 else 14 if level == 2 else 12, bold=True, color="1F4D78")
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_font(r, "宋体", 10)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_font(r, "宋体", 9.5, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        shade_cell(table.rows[0].cells[i], "F4F6F9")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], str(val), align=WD_ALIGN_PARAGRAPH.CENTER if len(str(val)) < 12 else WD_ALIGN_PARAGRAPH.LEFT)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    doc.add_paragraph()
    return table


def add_figure(doc, filename, caption, width=5.8):
    path = FIG_DIR / filename
    if path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width))
        add_caption(doc, caption)


def make_doc():
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.6)

    styles = doc.styles
    styles["Normal"].font.name = "宋体"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    styles["Normal"].font.size = Pt(11)

    # Cover
    for _ in range(2):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("硕士学位论文初稿")
    set_font(r, "黑体", 22, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("面向中国虚拟电厂的多模态语义增强深度强化学习动态优化与决策研究")
    set_font(r, "黑体", 18, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("A Study on Multimodal Semantic-Enhanced Deep Reinforcement Learning for Dynamic Optimization and Decision-Making of Chinese Virtual Power Plants")
    set_font(r, "Times New Roman", 12)
    doc.add_paragraph()
    meta = [
        ("学院", "数理学院"),
        ("专业", "大数据技术与工程"),
        ("研究方向", "虚拟电厂 / 智能调度 / 电力数据要素 / 大语言模型"),
        ("学生", "李星晨"),
        ("指导教师", "王继军 / 刘琼"),
        ("版本说明", "基于开题报告、开题PPT与IEEE Access稿件整理；中国数据实验为扩展方案，待实测补数。"),
    ]
    add_table(doc, ["项目", "内容"], meta, widths=[3.2, 11.5])
    doc.add_page_break()

    add_h(doc, "写作说明", 1)
    add_p(doc, "本稿不是学校最终排版模板，而是用于毕业论文撰写的可编辑内容初稿。文中已完成的实验结果来自现有 IEEE Access 稿件及其 reproducibility package；涉及中国数据的部分按毕业论文扩展需要设计为“新增实验方案”和“待实测结果位”。在正式提交前，应以导师确认的学校模板、真实运行结果和学校引用格式为准。")
    add_p(doc, "本稿的核心调整思路是：将 IEEE 论文中“电力数据资产估值与交易”的多模态数据、SAC 匹配和隐私实验基础，转化为“面向中国虚拟电厂运行场景的语义增强优化决策”毕业论文主线；同时把开题报告中的 BERT + SAC + VPP 仿真环境 + MILP 基准作为论文主体方法。")

    add_h(doc, "摘要", 1)
    add_p(doc, "在“双碳”目标和新型电力系统建设背景下，虚拟电厂通过聚合分布式光伏、储能、可控负荷和电力市场需求响应资源，成为提升新能源消纳能力和电力系统灵活性的重要技术形态。传统虚拟电厂优化调度方法主要依赖电价、负荷、光伏出力和储能状态等数值型时序数据，难以充分利用电力交易公告、气象灾害预警、调度指令和政策通知等非结构化文本信息，导致模型在极端天气、电价尖峰和调度约束突变场景下存在信息盲区。针对上述问题，本文提出一种面向中国虚拟电厂运行场景的多模态语义增强深度强化学习框架，将中文电力文本语义编码、虚拟电厂物理约束建模和 Soft Actor-Critic 连续控制策略相结合，构建“数值-语义”异构增强状态空间，实现虚拟电厂动态优化与可解释决策。")
    add_p(doc, "本文首先梳理虚拟电厂优化调度、深度强化学习、电力大模型和电力数据要素市场相关研究，明确现有方法在多源异构信息利用、约束满足和中国电力市场数据适配方面的不足。其次，设计面向中国数据的实验体系，建议接入国家能源局公开统计、区域电力交易中心披露信息、气象预警数据、分布式光伏和储能仿真数据，并将市场公告、天气预警和调度文本按时间戳对齐到 15 min 或 1 h 调度步长。然后，构建语言增强深度强化学习（LE-DRL）算法，利用 Chinese-BERT 或电力领域微调模型提取文本语义向量，通过特征适配层与负荷、电价、光伏、储能 SOC 等数值状态融合，并以经济收益、弃光惩罚、储能寿命损耗、购售电成本、越限惩罚和尾部风险为联合奖励函数。最后，设计规则策略、MILP、标准 SAC、文本消融 LE-DRL、多模态估值模型和隐私扰动实验等对比方案，以净收益、CVaR、最优性差距、约束违约率、调度响应时间和解释一致性为评价指标。")
    add_p(doc, "已有 IEEE 实验表明，多模态电力数据资产框架能够形成可复现实验基础：在 CAISO 2024 年 8777 个小时样本上，XGBoost 在单目标代理价值预测中取得最低 MAE，TransModal-ValueNet 在 MAE 上接近强基线但并未全面优于传统模型；在 20 微网 P2P 交易环境中，SAC 智能体相较随机策略取得约 3.41 倍累计奖励优势，成交率达到 90.13%；隐私实验显示，在满年度电价波动下，当前特征重要性引导的差分隐私预算分配不一定优于均匀分配。这些结果说明，多模态数据和强化学习在电力市场决策中具有应用潜力，但模型有效性必须通过中国区域市场数据和更多极端场景进一步验证。")
    add_p(doc, "关键词：虚拟电厂；深度强化学习；大语言模型；语义增强；Soft Actor-Critic；中国电力市场；多模态数据融合")

    add_h(doc, "Abstract", 1)
    add_p(doc, "Under China's dual-carbon strategy and the development of new power systems, virtual power plants aggregate distributed photovoltaics, energy storage, controllable loads and demand-response resources to improve renewable energy accommodation and system flexibility. Existing dispatching methods mainly rely on numerical time-series observations, such as load, price, renewable generation and state of charge, while non-structured information from market notices, weather warnings, dispatching instructions and policy documents is often ignored. This thesis proposes a multimodal semantic-enhanced deep reinforcement learning framework for Chinese virtual power plants. The framework embeds Chinese power-system texts into continuous semantic vectors, fuses them with physical numerical states, and trains a Soft Actor-Critic agent for dynamic dispatching under economic, risk and operational constraints. Existing IEEE experimental assets are reorganized into a graduate-thesis-oriented baseline, and a China-data experimental extension plan is provided for future empirical validation.")
    add_p(doc, "Keywords: virtual power plant; deep reinforcement learning; large language model; semantic enhancement; Soft Actor-Critic; Chinese power market; multimodal data fusion")
    doc.add_page_break()

    add_h(doc, "第1章 绪论", 1)
    add_h(doc, "1.1 研究背景", 2)
    for text in [
        "我国电力系统正经历从集中式化石能源主导向高比例新能源、多主体互动和电力市场化交易并存的结构性转型。国家能源局公开数据显示，2024 年我国全社会用电量达到 98521 亿千瓦时，同比增长 6.8%；截至 2024 年底，全国可再生能源装机达到 18.89 亿千瓦，其中风电 5.21 亿千瓦、太阳能发电 8.87 亿千瓦，可再生能源已成为电源结构变化最活跃的部分。这一趋势使电力系统面临更高比例随机性、更强不确定性和更复杂的时空耦合。虚拟电厂能够把分散的源、荷、储资源组织为可调度整体，是承接新能源消纳、需求响应和电力市场交易的重要载体。",
        "从调度决策角度看，虚拟电厂不仅需要处理负荷、电价、光伏出力、储能状态等结构化数据，还需要理解电力交易中心公告、气象灾害预警、需求响应通知、市场规则调整和调度指令等文本信息。例如，高温橙色预警往往意味着居民空调负荷上升和价格尖峰风险；新能源消纳通知可能意味着储能充放电策略需要提前调整；电力现货市场出清规则变化会影响套利空间。传统数值优化和多数深度强化学习方法无法直接理解这些语义信息，因此在突发事件下容易出现决策滞后。",
        "大语言模型和预训练语言模型为解决这一问题提供了新工具。BERT、Chinese-RoBERTa、MacBERT 等模型能够把中文文本映射为连续语义向量；深度强化学习能够在复杂环境中通过交互学习调度策略；SAC 通过最大熵目标在连续控制任务中兼顾收益和探索。若将中文电力文本语义编码与虚拟电厂数值状态融合，就可以构建信息更完备的增强状态空间，使智能体具备一定的事件感知和前瞻决策能力。"
    ]:
        add_p(doc, text)
    add_h(doc, "1.2 研究意义", 2)
    add_p(doc, "理论意义方面，本文把自然语言语义表示引入虚拟电厂马尔可夫决策过程，扩展了传统仅依赖数值状态的 MDP 建模方式。通过构造“数值状态 + 文本语义状态”的异构增强状态空间，可以从建模层面缓解部分可观测问题，并为多源异构信息驱动的电力系统智能控制提供可复用框架。")
    add_p(doc, "工程意义方面，本文面向中国电力市场和虚拟电厂场景，给出可执行的数据接入、文本事件构造、仿真环境搭建、算法训练和结果评价流程。相较单纯使用 CAISO 数据，增加中国数据有助于体现毕业论文与我国电力市场机制、双碳政策和新能源消纳实际需求的结合。")
    add_h(doc, "1.3 国内外研究现状", 2)
    add_p(doc, "虚拟电厂优化调度研究大致经历了数学规划、随机鲁棒优化、模型预测控制和数据驱动智能优化等阶段。传统 MILP、鲁棒优化和机会约束方法具有较强可解释性和约束表达能力，但对预测精度和模型参数依赖较强，在高维在线场景中求解压力较大。近年来，深度强化学习被用于储能套利、微网能量管理、虚拟电厂调峰辅助服务和多能流低碳调度。国内研究也开始关注注意力机制、约束强化学习和分层强化学习在综合能源系统中的应用。")
    add_p(doc, "大模型在电力行业的应用目前多集中于知识问答、设备故障诊断、辅助决策、预测解释和文本处理。已有综述指出，基于大语言模型的电力知识服务相对成熟，而多模态大模型、时序大模型以及大小模型协同的复杂决策场景仍处于探索阶段。由此可见，把电力文本语义编码直接嵌入虚拟电厂强化学习闭环，仍具有较明显的研究空间。")
    add_p(doc, "从已有 IEEE 稿件看，多模态 AI 可用于电力数据资产估值与交易，包括 TransModal-ValueNet、自然语言辅助 SAC 匹配和差分隐私效用测试。该工作为本论文提供了多模态数据处理、SAC 决策和隐私实验基础，但其目标偏向数据资产交易而非虚拟电厂运行调度。因此，毕业论文需要完成一次研究对象迁移：保留多模态数据、语义理解和强化学习方法，把实验核心转向中国虚拟电厂场景。")
    add_h(doc, "1.4 研究内容与创新点", 2)
    add_p(doc, "本文的主要研究内容包括：（1）建立面向中国虚拟电厂的多源异构数据体系；（2）构建融合中文文本语义和数值状态的 LE-DRL 决策框架；（3）设计包含规则策略、MILP、SAC、LE-DRL 和消融模型的实验体系；（4）复用并改造 IEEE 多模态电力数据实验基础，增加中国数据源、极端场景、尾部风险和可解释性实验；（5）形成适合硕士毕业论文的章节结构和可复现实验路线。")
    add_table(doc, ["创新点", "具体体现", "毕业论文中对应章节"], [
        ("状态空间创新", "将电力市场公告、气象预警、调度文本编码为语义向量，与负荷、电价、光伏、SOC 等数值状态拼接或交叉注意力融合", "第3章、第4章"),
        ("算法框架创新", "构建 Chinese-BERT/MacBERT + SAC 的语言增强强化学习框架，引入特征适配层和最大熵连续控制", "第4章"),
        ("实验体系创新", "在 IEEE 现有多模态实验基础上增加中国区域电力市场、气象预警、新能源消纳和虚拟电厂运行场景", "第5章、第6章"),
        ("风险评价创新", "在收益指标外加入 CVaR、约束违约率、储能寿命损耗、弃光惩罚和极端事件响应能力", "第5章、第6章"),
    ], widths=[3.2, 8.0, 4.2])

    add_h(doc, "第2章 理论基础与关键技术", 1)
    add_h(doc, "2.1 虚拟电厂优化调度模型", 2)
    add_p(doc, "虚拟电厂由分布式电源、储能系统、可控负荷和市场接口组成。设调度步长为 Δt，虚拟电厂在时刻 t 的数值状态可表示为 s_t^num=[P_pv,t, P_w,t, P_load,t, SOC_t, λ_buy,t, λ_sell,t, ...]。动作 a_t 通常表示储能充放电功率、可控负荷调整量和购售电决策。")
    add_p(doc, "储能状态转移可表示为：SOC_{t+1}=SOC_t+η_ch P_ch,t Δt/E_cap-(1/η_dis) P_dis,t Δt/E_cap。调度目标是在满足 SOC 边界、功率边界、能量平衡和市场交易规则的条件下最大化累计净收益，并降低极端场景下的损失风险。")
    add_h(doc, "2.2 Soft Actor-Critic 算法", 2)
    add_p(doc, "SAC 属于离策略 Actor-Critic 方法，适合连续动作空间。其最大熵目标为 J(π)=Σ E[r(s_t,a_t)+αH(π(·|s_t))]，其中 H 为策略熵，α 为温度系数。与 DDPG、TD3 相比，SAC 通过熵正则鼓励探索，在高不确定性电价和新能源出力场景中更有利于避免策略过早收敛到局部最优。")
    add_h(doc, "2.3 中文电力文本语义编码", 2)
    add_p(doc, "电力市场公告、天气预警和调度通知通常具有强领域语义。例如“高温红色预警”“需求响应邀约”“新能源消纳压力增大”“日前价格异常波动”等短文本会改变虚拟电厂对未来负荷、电价和储能备用的判断。本文建议采用 Chinese-BERT、MacBERT 或电力语料微调模型，将文本 T_t 编码为 h_t^sem=BERT(T_t)，再通过 MLP、注意力或门控机制降维到 d_sem。")
    add_h(doc, "2.4 多模态融合与数据资产估值基础", 2)
    add_p(doc, "IEEE 稿件中的 TransModal-ValueNet 使用跨模态注意力将时间、负荷、市场价格、可再生能源和气象变量映射到代理价值评分。虽然该任务与 VPP 调度不同，但其多模态投影层、[CLS] token、cross-attention 和自监督重构训练可迁移到虚拟电厂状态表征。毕业论文可将该模块定位为“辅助表征学习模块”，用于提取电力数据质量、波动性和风险信号，而不必把数据资产估值作为主贡献。")

    add_h(doc, "第3章 面向中国数据的多源异构数据体系构建", 1)
    add_h(doc, "3.1 数据源选择原则", 2)
    add_p(doc, "中国数据扩展不宜只把 CAISO 替换为单一公开表格，而应构建“公开统计 + 区域市场 + 气象预警 + 仿真资源”的组合数据集。原因在于国内真实虚拟电厂分钟级运行数据和交易明细通常不可公开获取，硕士论文更可行的路线是以公开官方数据校准宏观趋势，以区域电力交易披露和气象数据构造时序场景，再用可解释的仿真模型生成 VPP 内部源荷储资源。")
    add_table(doc, ["数据类别", "建议中国数据源", "用途", "采样粒度建议"], [
        ("用电与电源结构", "国家能源局全社会用电量、全国电力工业统计、可再生能源并网运行情况", "校准负荷增长、新能源占比和实验背景", "月度/年度"),
        ("区域市场价格", "广东、山西、山东等电力交易中心公开披露的现货市场公告、交易简报和出清信息", "构造日前/实时电价与价格尖峰场景", "日度/小时/15min，按可获得性"),
        ("气象与预警", "中国气象数据网、气象预警公开接口、Open-Meteo 中国坐标历史天气、和风天气等", "温度、风速、辐照、降水、极端天气事件", "小时/日度"),
        ("新能源出力", "国家能源局、全国新能源消纳监测预警中心、区域公开新能源运行数据；必要时用辐照/风速模型生成", "光伏、风电出力与弃电风险", "小时/15min"),
        ("文本事件", "电力交易公告、气象预警、需求响应通知、政策新闻、调度规则文本", "构造语义增强状态 h_t^sem", "事件触发，对齐到调度步长"),
        ("VPP内部资源", "基于中国典型工商业负荷、分布式光伏、储能参数和可控负荷比例仿真", "形成可交互 Gymnasium 环境", "15min/1h"),
    ], widths=[2.7, 5.7, 5.1, 2.2])
    add_h(doc, "3.2 数据预处理与时间对齐", 2)
    add_p(doc, "数值型数据应统一到 15 min 或 1 h 调度步长。对于低频统计数据，可用于宏观参数校准而不直接进入 RL 状态；对于气象和市场价格，应进行缺失值填补、异常值截尾、节假日特征构造和标准化。文本事件采用事件触发衰减机制：若某公告在 t0 发布，则其语义影响在后续 k 个调度步内按 exp(-β(t-t0)) 衰减。这样可以避免文本事件只在一个时间点生效。")
    add_p(doc, "建议构建三套中国实验场景：常规日场景、高温/寒潮极端负荷场景、市场价格尖峰与新能源消纳受限场景。每套场景至少包含负荷、光伏、储能 SOC、购售电价格、天气变量和文本语义状态。")
    add_h(doc, "3.3 文本事件标签体系", 2)
    add_table(doc, ["文本类型", "示例", "语义标签", "预期调度影响"], [
        ("气象预警", "高温橙色预警、寒潮蓝色预警、大风预警", "负荷上升、光伏/风电波动、风险等级", "提前保留储能、增加备用容量"),
        ("市场公告", "日前价格异常、现货结算试运行、需求响应邀约", "价格趋势、交易机会、约束变化", "调整购售电策略和报价区间"),
        ("调度通知", "削峰填谷、负荷管理、新能源消纳压力", "调峰需求、可调负荷、弃电风险", "提升响应动作、降低弃光弃风"),
        ("政策规则", "绿电交易、辅助服务、现货市场规则修订", "收益机制变化、交易边界", "改变奖励函数参数或约束边界"),
    ], widths=[2.5, 4.4, 4.0, 5.0])

    add_h(doc, "第4章 语言增强深度强化学习方法设计", 1)
    add_h(doc, "4.1 总体框架", 2)
    add_p(doc, "本文提出的 LE-DRL 框架由四层组成：数据层、语义编码层、强化学习决策层和解释评估层。数据层负责同步中国区域电力市场、气象、负荷、光伏和文本事件；语义编码层使用中文预训练模型输出文本语义向量；决策层使用 SAC 输出连续充放电或购售电动作；解释评估层将动作、状态和文本事件关联，生成可读的调度原因。")
    add_figure(doc, "system_architecture.png", "图4-1 IEEE 现有多模态电力数据框架，可改造为本文 LE-DRL 数据与决策架构参考图", width=5.5)
    add_h(doc, "4.2 异构增强状态空间", 2)
    add_p(doc, "增强状态定义为 s_t^aug = [s_t^num, z_t^sem, z_t^risk]。其中 s_t^num 包含负荷、光伏、风电、SOC、电价、天气数值变量；z_t^sem 为中文文本编码后的语义向量；z_t^risk 为由文本分类器或规则模型得到的风险标签，例如高温负荷风险、价格尖峰风险和新能源消纳风险。")
    add_p(doc, "为避免高维语义向量压制低维数值状态，本文建议在输入 SAC 网络前增加 Feature Adapter：z_t=σ(W_g[h_t^sem;s_t^num])⊙W_s h_t^sem，其中 σ 表示门控函数。该结构能够让模型在文本无明显事件时降低语义权重，在极端预警或市场公告出现时提高语义权重。")
    add_h(doc, "4.3 奖励函数设计", 2)
    add_p(doc, "奖励函数应服务于毕业论文的实验解释。建议定义 r_t = R_sell,t - C_buy,t - C_bat,t - C_curt,t - C_violate,t - ρ·Risk_t。其中 R_sell,t 为售电收益，C_buy,t 为购电成本，C_bat,t 为储能寿命损耗，C_curt,t 为弃光弃风惩罚，C_violate,t 为 SOC 或功率越限惩罚，Risk_t 为尾部损失或价格尖峰暴露。")
    add_h(doc, "4.4 MILP 基准模型", 2)
    add_p(doc, "为避免强化学习结果缺乏理论参照，建议建立滚动 MILP 基准。MILP 在已知未来 24 h 负荷、光伏和电价预测的条件下求解储能充放电和购售电策略，可作为局部最优或理论上界。毕业论文中应报告 LE-DRL 相对 MILP 的最优性差距，而不是只与随机策略比较。")
    add_h(doc, "4.5 可解释性模块", 2)
    add_p(doc, "可解释性不必依赖大型闭源 LLM。可采用“规则模板 + 文本事件 + 状态贡献”的方式生成解释。例如：当高温预警出现、负荷预测上升且电价处于上分位区间时，解释为“模型选择降低放电幅度以保留晚高峰套利和削峰能力”。若使用大语言模型生成自然语言解释，应把输入限制为状态摘要和动作摘要，避免凭空生成不存在的数据。")

    add_h(doc, "第5章 实验设计与中国数据扩展方案", 1)
    add_h(doc, "5.1 IEEE 已有实验基础", 2)
    add_p(doc, "现有 IEEE 包提供了可复用的实验基础：CAISO 2024 年全年度市场数据、Open-Meteo 气象变量、P2P 微网交易环境、多模态估值模型、SAC 交易策略、差分隐私扰动实验和完整图表。毕业论文可在第6章中把这些结果作为“前期实验基础”，同时强调其地理区域是美国加州，因此不能直接代表中国虚拟电厂运行效果。")
    add_table(doc, ["实验模块", "IEEE已有结果", "在毕业论文中的用途"], [
        ("多模态估值", "8777 小时样本；XGBoost MAE=0.0297，TransModal-ValueNet MAE=0.0313", "证明多模态电力数据处理流程可运行；作为中国数据替换前的基础结果"),
        ("SAC交易", "SAC 平均奖励 13436，总奖励 46.43M，较随机策略约 3.41 倍", "证明 SAC 在电力交易/匹配类连续决策中有效，可迁移到 VPP 储能调度"),
        ("隐私实验", "ε=10 时均匀 DP 保留 25.9% 非隐私价值，高于当前自适应分配 16.9%", "作为负结果讨论，说明数据要素交易与隐私保护需要验证驱动"),
        ("消融与鲁棒性", "去除市场、发电、天气等模态后性能变化不同", "支撑多源异构数据的重要性"),
    ], widths=[3.0, 6.5, 6.5])
    add_figure(doc, "baseline_comparison_bar.png", "图5-1 IEEE 前期实验：多模型代理价值预测结果", width=5.2)
    add_figure(doc, "drl_cumulative_reward.png", "图5-2 IEEE 前期实验：SAC 与随机策略累计奖励对比", width=5.2)
    add_h(doc, "5.2 建议新增中国数据实验", 2)
    add_table(doc, ["实验编号", "实验名称", "目的", "关键对比"], [
        ("E1", "中国区域常规日 VPP 调度", "验证 LE-DRL 在普通负荷与价格波动下的经济性", "Rule、MILP、SAC、LE-DRL"),
        ("E2", "高温/寒潮文本预警实验", "验证文本语义是否能提升提前储能和负荷响应能力", "SAC 无文本 vs LE-DRL 有文本"),
        ("E3", "价格尖峰与需求响应公告实验", "测试市场公告对购售电策略的影响", "无公告、规则公告、BERT语义公告"),
        ("E4", "新能源消纳受限实验", "评估弃光惩罚和消纳文本对策略的影响", "不同弃电惩罚系数、不同文本衰减系数"),
        ("E5", "中国数据多模态消融", "衡量负荷、天气、价格、文本、光伏各模态贡献", "去除单一模态、去除文本、只用数值"),
        ("E6", "可解释性与鲁棒性实验", "评估解释一致性、随机种子稳定性和极端事件尾部风险", "多随机种子、CVaR、约束违约率"),
    ], widths=[1.7, 4.0, 5.8, 4.9])
    add_h(doc, "5.3 评价指标", 2)
    add_p(doc, "经济性指标包括总净收益、日均收益、购电成本、售电收益和相对规则基线提升率。安全性指标包括 SOC 越限率、功率越限率、电力平衡误差和储能循环损耗。风险指标包括收益 CVaR、极端价格日最差收益和最大回撤。学习指标包括收敛轮数、平均奖励、策略熵和训练稳定性。文本指标包括事件召回率、解释一致性和文本消融收益差异。")
    add_h(doc, "5.4 实验复现流程", 2)
    add_p(doc, "建议代码目录按 data/raw_china、data/processed、envs/vpp_env.py、models/semantic_encoder.py、agents/sac_ledrl.py、experiments、figures 和 thesis_tables 组织。每次实验保存 config.yaml、随机种子、训练日志、模型权重和结果 CSV。毕业论文只放关键表图，完整日志作为附录或代码仓库材料。")

    add_h(doc, "第6章 结果分析与讨论", 1)
    add_h(doc, "6.1 前期结果：多模态电力数据实验", 2)
    add_p(doc, "IEEE 前期实验显示，多模态电力数据并不必然让深度模型全面超过传统机器学习基线。在线性关系较强的单目标代理价值预测中，线性回归和 XGBoost 仍然具有竞争力。这一点对毕业论文很重要：论文不应简单宣称“大模型一定更好”，而应把创新点放在文本语义进入决策闭环后对极端事件响应、风险控制和可解释性的改善。")
    add_figure(doc, "ablation_results.png", "图6-1 IEEE 前期实验：多模态消融结果", width=5.4)
    add_figure(doc, "privacy_utility_tradeoff.png", "图6-2 IEEE 前期实验：差分隐私效用权衡", width=5.4)
    add_h(doc, "6.2 中国数据实验结果表格模板", 2)
    add_p(doc, "以下表格为中国数据实测后应补入的核心结果格式。当前不填虚构数值，只保留“待实测”。")
    add_table(doc, ["模型", "总净收益", "相对规则提升", "CVaR(5%)", "SOC越限率", "平均响应时间", "说明"], [
        ("Rule-based", "待实测", "0%", "待实测", "待实测", "待实测", "基础规则策略"),
        ("MILP", "待实测", "理论上界/滚动上界", "待实测", "待实测", "待实测", "需已知预测窗口"),
        ("SAC", "待实测", "待实测", "待实测", "待实测", "待实测", "仅数值状态"),
        ("LE-DRL", "待实测", "待实测", "待实测", "待实测", "待实测", "数值+文本语义"),
        ("LE-DRL w/o Text", "待实测", "待实测", "待实测", "待实测", "待实测", "文本消融"),
    ], widths=[2.8, 2.3, 2.5, 2.2, 2.2, 2.2, 2.7])
    add_h(doc, "6.3 预期分析逻辑", 2)
    add_p(doc, "若 LE-DRL 在常规场景收益提升不明显，但在高温预警、价格尖峰或需求响应场景中显著降低尾部损失，则论文结论应表述为“语义增强主要改善极端事件和信息不完备场景下的风险控制”，而不是泛化为所有场景收益最高。若 LE-DRL 在收益上超过 SAC 但越限率提高，需要重新调整奖励函数和动作投影层，优先保证安全约束。")
    add_p(doc, "若中文文本编码模型效果不稳定，可比较三种方案：关键词规则、Chinese-BERT/MacBERT 句向量、领域微调模型。对于硕士论文，能够证明“文本事件加入后在若干典型场景中改变策略并改善风险指标”即可形成较完整贡献。")

    add_h(doc, "第7章 结论与展望", 1)
    add_p(doc, "本文围绕中国虚拟电厂动态优化与决策问题，提出了多模态语义增强深度强化学习研究框架。该框架把中文电力文本、气象预警、市场公告和数值型源荷储状态统一到增强状态空间，并使用 SAC 学习连续调度策略。与传统仅依赖数值时序的强化学习相比，该框架的核心价值在于利用文本事件补充未来风险信息，使虚拟电厂在极端天气、电价尖峰和调度规则变化场景下具备更强的前瞻性。")
    add_p(doc, "现有 IEEE 实验为本文提供了多模态电力数据处理、SAC 决策和隐私效用测试基础，但其数据源主要来自 CAISO 和 Open-Meteo，不足以支撑“中国虚拟电厂”结论。因此，后续工作应优先完成中国区域数据采集、文本事件对齐、LE-DRL 训练和消融实验，把“建议实验”转化为真实结果。")
    add_p(doc, "未来可进一步研究三方面内容：一是引入多智能体强化学习，刻画多个虚拟电厂或聚合商之间的竞价互动；二是结合联邦学习和差分隐私，实现跨主体数据协同建模；三是训练面向电力调度语料的领域语言模型，提高公告、预警和规则文本的语义识别能力。")

    add_h(doc, "参考文献", 1)
    refs = [
        "[1] 国家能源局. 2024年全社会用电量同比增长6.8%[EB/OL]. 2025-01-20.",
        "[2] 国家能源局. 2024年可再生能源并网运行情况[EB/OL]. 2025-01-27.",
        "[3] 国家能源局. 国家能源局发布2024年全国电力工业统计数据[EB/OL]. 2025-01-21.",
        "[4] 国家能源局. 《关于做好新能源消纳工作保障新能源高质量发展的通知》政策解读[EB/OL]. 2024-06-04.",
        "[5] 广州电力交易中心. 南方区域电力市场连续结算试运行半年记[EB/OL]. 2026-01-06.",
        "[6] 艾欣, 庞博, 孙毅, 等. 虚拟电厂多时间尺度优化调度研究综述[J]. 电网技术, 2024.",
        "[7] 陈来军, 孙文, 周博, 等. 基于SAC深度强化学习的虚拟电厂调峰辅助服务优化[J]. 华北电力大学学报(自然科学版), 2024.",
        "[8] 王雨晴, 艾欣, 孙毅, 等. 大语言模型在电力系统中的应用综述与展望[J]. 电网技术, 2024.",
        "[9] 融合注意力机制与SAC算法的虚拟电厂多能流低碳调度[J]. 电力工程技术.",
        "[10] 基于分层约束强化学习的综合能源多微网系统优化调度[J]. 电工技术学报, 2024.",
        "[11] Devlin J, Chang M W, Lee K, et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding[C]. NAACL, 2019.",
        "[12] Haarnoja T, Zhou A, Abbeel P, Levine S. Soft Actor-Critic: Off-Policy Maximum Entropy Deep Reinforcement Learning with a Stochastic Actor[C]. ICML, 2018.",
        "[13] Vaswani A, et al. Attention is All You Need[C]. NeurIPS, 2017.",
        "[14] Raffin A, et al. Stable-Baselines3: Reliable Reinforcement Learning Implementations[J]. JMLR, 2021.",
        "[15] Dwork C, Roth A. The Algorithmic Foundations of Differential Privacy[J]. Foundations and Trends in Theoretical Computer Science, 2014.",
    ]
    for ref in refs:
        add_p(doc, ref, first_line=False)

    add_h(doc, "附录A：论文从开题与IEEE稿件扩展到毕业论文的执行清单", 1)
    add_table(doc, ["任务", "当前状态", "下一步"], [
        ("开题报告主线", "已明确为 LLM/DRL 融合 VPP 动态优化", "继续保留，作为论文核心问题"),
        ("IEEE稿件", "电力数据资产估值与交易，有完整图表和实验", "作为前期多模态与SAC基础，改写成章节支撑"),
        ("中国数据", "尚未完成真实采集和训练", "优先选定 1 个省级市场 + 气象 + 仿真 VPP"),
        ("新增实验", "已设计 E1-E6", "按重要性先跑 E1、E2、E5"),
        ("文献", "已有中英文混合参考", "正式版需补充 CNKI/万方可查的中文文献并统一格式"),
        ("学校模板", "本稿未套用学校模板", "拿到模板后迁移标题页、目录、页眉页脚和参考文献格式"),
    ], widths=[3.8, 5.2, 6.2])

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("硕士学位论文初稿 - 中国数据扩展版")
    set_font(r, "宋体", 9, color="555555")

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    out = make_doc()
    print(out)
