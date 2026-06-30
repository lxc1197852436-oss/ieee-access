from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
MD_PATH = ROOT / "reports" / "full_thesis_draft.md"
OUT_PATH = ROOT / "outputs" / "上海电力大学_硕士论文初稿_整合版.docx"


def set_east_asia(run, font_name: str) -> None:
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_style_font(style, font_name: str, size_pt: int, bold: bool = False, color: str = "000000") -> None:
    style.font.name = font_name
    style._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(text) <= 12 else WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text.strip())
    set_east_asia(r, "宋体")
    r.font.size = Pt(9)
    r.font.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code_paragraph(doc: Document, line: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(line)
    set_east_asia(r, "Courier New")
    r.font.name = "Courier New"
    r.font.size = Pt(8.5)


def add_markdown_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    parts = re.split(r"(`[^`]+`|\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = p.add_run(part[1:-1])
            set_east_asia(run, "Courier New")
            run.font.name = "Courier New"
            run.font.size = Pt(10)
        elif part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            set_east_asia(run, "宋体")
            run.font.bold = True
        else:
            run = p.add_run(part)
            set_east_asia(run, "宋体")


def add_image(doc: Document, alt: str, image_path: str) -> None:
    path = Path(image_path)
    if not path.is_absolute():
        path = (ROOT / "reports" / path).resolve()
    if not path.exists():
        add_markdown_paragraph(doc, f"[图片缺失：{alt}，路径：{image_path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.2))
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = caption.add_run(alt)
    set_east_asia(r, "宋体")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor.from_string("4A5568")


def split_table_row(line: str) -> list[str]:
    line = line.strip()
    if line.startswith("|"):
        line = line[1:]
    if line.endswith("|"):
        line = line[:-1]
    return [cell.strip() for cell in line.split("|")]


def is_table_separator(line: str) -> bool:
    cells = split_table_row(line)
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell.strip()) for cell in cells if cell.strip())


def add_table(doc: Document, rows: list[str]) -> None:
    if len(rows) < 2:
        return
    header = split_table_row(rows[0])
    body = [split_table_row(r) for r in rows[2:] if r.strip()]
    table = doc.add_table(rows=1, cols=len(header))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, h in enumerate(header):
        set_cell_text(table.rows[0].cells[i], h, bold=True)
        set_cell_shading(table.rows[0].cells[i], "F2F4F7")
    for row in body:
        cells = table.add_row().cells
        for i in range(len(header)):
            set_cell_text(cells[i], row[i] if i < len(row) else "")
    doc.add_paragraph()


def build_docx() -> Path:
    text = MD_PATH.read_text(encoding="utf-8")
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    styles = doc.styles
    set_style_font(styles["Normal"], "宋体", 11)
    set_style_font(styles["Heading 1"], "黑体", 16, bold=True)
    set_style_font(styles["Heading 2"], "黑体", 14, bold=True)
    set_style_font(styles["Heading 3"], "黑体", 12, bold=True)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("上海电力大学硕士学位论文初稿")

    lines = text.splitlines()
    i = 0
    in_code = False
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith("```"):
            in_code = not in_code
            i += 1
            continue
        if in_code:
            add_code_paragraph(doc, line)
            i += 1
            continue
        if not line.strip():
            i += 1
            continue
        image_match = re.match(r"^!\[(.*?)\]\((.*?)\)\s*$", line)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_rows.append(lines[i])
                i += 1
            add_table(doc, table_rows)
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:].strip(), level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if line.startswith("# 大语言模型") else WD_ALIGN_PARAGRAPH.LEFT
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue
        if re.match(r"^[-*] ", line):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(0.74)
            r = p.add_run("• " + line[2:].strip())
            set_east_asia(r, "宋体")
            i += 1
            continue
        if re.match(r"^\d+\. ", line):
            p = doc.add_paragraph(style=None)
            p.paragraph_format.left_indent = Cm(0.74)
            r = p.add_run(line.strip())
            set_east_asia(r, "宋体")
            i += 1
            continue
        add_markdown_paragraph(doc, line)
        i += 1

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_docx())
